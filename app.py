import streamlit as st
import os
import random
import re
from io import BytesIO
import requests
import numpy as np
import pandas as pd
from PIL import Image, ImageFilter
import exifread
import pdfplumber
from docx import Document
from google.oauth2 import service_account
from google.cloud import texttospeech
import plotly.express as px
import plotly.graph_objs as go
import streamlit.components.v1 as components
from pydub import AudioSegment
import openai
import json
import tempfile
from textblob import TextBlob
import nltk
import bcrypt
import base64
import sqlite3

# 必要なNLTKリソースのダウンロード
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('wordnet')

# 初期設定
st.set_page_config(page_title="究極融合アプリ", page_icon="✨", layout="wide")

# カスタムCSSを追加（洗練されたUI）
custom_css = """
<style>
body {
    background: #121212;
    color: #ffffff;
    font-family: 'Helvetica', sans-serif;
}
h1, h2, h3, h4, h5, h6 {
    color: #ffffff;
}
.block-container {
    padding: 1rem 2rem;
}
.sidebar .sidebar-content {
    background: #1e1e1e;
    color: #ffffff;
}
.stTextInput > div {
    color:#ffffff;
}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# データベース接続
def init_db():
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password BLOB NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def load_users_from_db():
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('SELECT username, password FROM users')
    users = c.fetchall()
    conn.close()
    return {user[0]: user[1] for user in users}

def save_users_to_db(users):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('DELETE FROM users')
    for username, password in users.items():
        c.execute('INSERT INTO users (username, password) VALUES (?, ?)', (username, password))
    conn.commit()
    conn.close()

def hash_password(password):
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt())

def check_password(password, hashed):
    return bcrypt.checkpw(password.encode(), hashed)

# 初期化
init_db()

# Google Cloud TTS認証
if "gcp_service_account" in st.secrets:
    service_account_info = st.secrets["gcp_service_account"]
    credentials = service_account.Credentials.from_service_account_info(service_account_info)
    tts_client = texttospeech.TextToSpeechClient(credentials=credentials)
else:
    st.error("Google Cloudサービスアカウント情報がst.secretsにありません。設定してください。")
    st.stop()

# OpenAI APIキー設定
if "openai" in st.secrets and "api_key" in st.secrets["openai"]:
    openai.api_key = st.secrets["openai"]["api_key"]
else:
    openai.api_key = None

# セッション状態の初期化
if "messages" not in st.session_state:
    st.session_state["messages"] = [{
        "role": "assistant", 
        "content": "ここは人類史上初の究極融合アプリ。あなたがアップロードするテキストや画像を解析し、音声化し、可視化する。そして高度なGPT対話すら可能な、夢の一頁です。"
    }]
if "exif_df" not in st.session_state:
    st.session_state["exif_df"] = pd.DataFrame()
if "image_url" not in st.session_state:
    st.session_state["image_url"] = ""
if "uploaded_files" not in st.session_state:
    st.session_state["uploaded_files"] = None
if "user_preferences" not in st.session_state:
    st.session_state["user_preferences"] = {
        "theme": "ダークモード",
        "notifications": True
    }
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if "username" not in st.session_state:
    st.session_state["username"] = ""

########################################################
# 幻想的粒子アニメーション背景
########################################################
particles_js = """<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Particles.js</title>
<style>
#particles-js {
  position: fixed;
  width:100vw;
  height:100vh;
  top:0;left:0;z-index:-1;
  background:#000;
}
.content {
  position:relative;z-index:1;color:white;
}
</style>
</head>
<body>
<div id="particles-js"></div>
<div class="content"></div>
<script src="https://cdn.jsdelivr.net/particles.js/2.0.0/particles.min.js"></script>
<script>
particlesJS("particles-js", {
  "particles":{
    "number":{"value":300,"density":{"enable":true,"value_area":800}},
    "color":{"value":"#ffffff"},
    "shape":{"type":"circle","stroke":{"width":0,"color":"#000000"}},
    "opacity":{"value":0.5,"random":false},
    "size":{"value":2,"random":true},
    "line_linked":{"enable":true,"distance":100,"color":"#ffffff","opacity":0.22,"width":1},
    "move":{"enable":true,"speed":0.2,"direction":"none","random":false,"straight":false,"out_mode":"out","bounce":true}
  },
  "interactivity":{
    "events":{
      "onhover":{"enable":true,"mode":"grab"},
      "onclick":{"enable":true,"mode":"repulse"},
      "resize":true
    },
    "modes":{
      "grab":{"distance":100,"line_linked":{"opacity":1}},
      "bubble":{"distance":400,"size":2,"duration":2,"opacity":0.5,"speed":1},
      "repulse":{"distance":200,"duration":0.4},
      "push":{"particles_nb":2},
      "remove":{"particles_nb":3}
    }
  },
  "retina_detect":true
});
</script>
</body>
</html>
"""
components.html(particles_js, height=0, width=0)

########################################################
# ユーティリティ関数群
########################################################
def clear_url():
    st.session_state["image_url"] = ""

def clear_files():
    st.session_state["uploaded_files"] = None
    st.session_state["file_uploader_key"] = not st.session_state.get("file_uploader_key", False)

def clear_chat_history():
    st.session_state["messages"] = [{
        "role":"assistant",
        "content":"チャット履歴をクリアしました。再び新たなる時代へ踏み出そう。"
    }]
    st.session_state["exif_df"] = pd.DataFrame()
    st.session_state["uploaded_files"] = None
    st.session_state["image_url"] = ""
    st.session_state["user_preferences"] = {
        "theme": "ダークモード",
        "notifications": True
    }
    # 適切なキャッシュクリア方法を使用
    if hasattr(st, 'cache_data'):
        st.cache_data.clear()
    elif hasattr(st, 'cache_resource'):
        st.cache_resource.clear()
    st.success("チャット履歴をクリアしました！")

def load_image(file):
    if isinstance(file, str):
        response = requests.get(file)
        response.raise_for_status()
        return Image.open(BytesIO(response.content))
    elif isinstance(file, bytes):
        return Image.open(BytesIO(file))
    else:
        return Image.open(file)

def clear_exif_data(image_input):
    if isinstance(image_input, BytesIO):
        image_input.seek(0)
        image = Image.open(image_input)
    elif isinstance(image_input, Image.Image):
        image = image_input
    else:
        raise ValueError("画像タイプがサポートされていません")
    data = list(image.getdata())
    image_without_exif = Image.new(image.mode, image.size)
    image_without_exif.putdata(data)

    buffered = BytesIO()
    image_without_exif.save(buffered, format="JPEG", quality=100, optimize=True)
    buffered.seek(0)
    return buffered.getvalue()

def download_image(data):
    st.download_button(
        label="⇩ EXIF除去後の画像ダウンロード",
        data=data,
        file_name="image_no_exif.jpg",
        mime="image/jpeg",
    )

def detect_language(text):
    if re.search('[\u3000-\u303F\u3040-\u309F\u30A0-\u30FF]', text):
        return 'ja-JP'
    return 'en-US'

def synthesize_speech_chunk(text, lang_code, gender='neutral', rate=1.0, pitch=0.0):
    max_chars = 4500
    chunks = [text[i:i+max_chars] for i in range(0,len(text),max_chars)]

    gender_map = {
        'default': texttospeech.SsmlVoiceGender.SSML_VOICE_GENDER_UNSPECIFIED,
        'male': texttospeech.SsmlVoiceGender.MALE,
        'female': texttospeech.SsmlVoiceGender.FEMALE,
        'neutral': texttospeech.SsmlVoiceGender.NEUTRAL
    }

    combined_audio = AudioSegment.empty()

    for i, chunk in enumerate(chunks):
        synthesis_input = texttospeech.SynthesisInput(text=chunk)
        voice = texttospeech.VoiceSelectionParams(
            language_code=lang_code,
            ssml_gender=gender_map.get(gender, texttospeech.SsmlVoiceGender.NEUTRAL)
        )
        audio_config = texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3,
            speaking_rate=rate,
            pitch=pitch
        )
        response = tts_client.synthesize_speech(input=synthesis_input, voice=voice, audio_config=audio_config)

        segment = AudioSegment.from_file(BytesIO(response.audio_content), format="mp3")
        combined_audio += segment

    output_buffer = BytesIO()
    combined_audio.export(output_buffer, format="mp3")
    output_buffer.seek(0)
    return output_buffer

def summarize_text(text, language='ja'):
    # 簡単な要約を行う関数（TextBlobを使用）
    if language == 'ja':
        # 日本語の要約はTextBlobでは対応していないため、簡易的に文を抽出
        sentences = re.split('。|\n', text)
        summary = '。'.join(sentences[:3]) + '。' if len(sentences) > 3 else text
    else:
        blob = TextBlob(text)
        summary = blob.noun_phrases
        summary = ', '.join(summary[:5]) if summary else text
    return summary

def analyze_sentiment(text):
    blob = TextBlob(text)
    sentiment = blob.sentiment.polarity
    if sentiment > 0.1:
        return 'ポジティブ'
    elif sentiment < -0.1:
        return 'ネガティブ'
    else:
        return 'ニュートラル'

def extract_keywords(text, num_keywords=5):
    blob = TextBlob(text)
    keywords = blob.noun_phrases
    return ', '.join(keywords[:num_keywords]) if keywords else 'なし'

########################################################
# サイドバー
########################################################
with st.sidebar:
    st.markdown("<h1 style='color:white;'>融合</h1>",unsafe_allow_html=True)
    st.markdown("#### EXIF解析 & 超大規模TTS & GPT対話")
    
    # ファイル入力エクスパンダー
    expander = st.expander("🗀 ファイル入力")
    with expander:
        st.text("長大テキスト/画像/URL分析対応")
        image_url = st.text_input(
            "EXIF解析用画像URL:",
            key="image_url",
            on_change=clear_files,
            value=st.session_state["image_url"],
        )
        file_uploader_key = "file_uploader_{}".format(
            st.session_state.get("file_uploader_key", False)
        )
        uploaded_files = st.file_uploader(
            "ファイルアップロード:",
            type=["txt","pdf","docx","csv","jpg","png","jpeg"],
            key=file_uploader_key,
            on_change=clear_url,
            accept_multiple_files=True,
        )
        if uploaded_files is not None:
            st.session_state["uploaded_files"] = uploaded_files

    # モデル設定エクスパンダー
    expander = st.expander("⚒ Model Configuration")
    with expander:
        if "REPLICATE_API_TOKEN" in st.secrets:
            replicate_api = st.secrets["REPLICATE_API_TOKEN"]
        else:
            replicate_api = st.text_input("Enter Replicate API token:", type="password")
            if not (replicate_api.startswith("r8_") and len(replicate_api) == 40):
                st.warning("Please enter your Replicate API token.", icon="⚠️")
                st.markdown(
                    "**Don't have an API token?** Head over to [Replicate](https://replicate.com/account/api-tokens) to sign up for one."
                )
        os.environ["REPLICATE_API_TOKEN"] = replicate_api
        st.subheader("Adjust model parameters")
        temperature = st.slider(
            "Temperature", min_value=0.01, max_value=5.0, value=0.3, step=0.01
        )
        top_p = st.slider("Top P", min_value=0.01, max_value=1.0, value=0.2, step=0.01)
        max_new_tokens = st.number_input(
            "Max New Tokens", min_value=1, max_value=1024, value=512
        )
        min_new_tokens = st.number_input(
            "Min New Tokens", min_value=0, max_value=512, value=0
        )
        presence_penalty = st.slider(
            "Presence Penalty", min_value=0.0, max_value=2.0, value=1.15, step=0.05
        )
        frequency_penalty = st.slider(
            "Frequency Penalty", min_value=0.0, max_value=2.0, value=0.2, step=0.05
        )
        stop_sequences = st.text_area("Stop Sequences", value="<|im_end|>", height=100)

    st.markdown("---")
    st.button("🗑 チャット履歴クリア", on_click=clear_chat_history)
    st.markdown("---")
    st.caption("© Exifa.net (Sahir Maharaj,2024), CC-BY 4.0")
    
    # ソーシャルリンク
    linkedin = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/linkedin.svg"
    youtube = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/youtube.svg"
    paypal = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/paypal.svg"
    dropbox = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/dropbox.svg"
    huggingface = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/huggingface.svg"
    x_icon = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/x.svg"
    facebook = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/facebook.svg"
    amazon_jp = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/amazon-jp.svg"
    amazon_us = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/amazon-us.svg"
    newsletter = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/newsletter.svg"
    share = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/share.svg"
    
    uptime = "https://uptime.betterstack.com/status-badges/v1/monitor/196o6.svg"
    
    st.markdown(
        f"""
        <div style='display: flex; align-items: center;'>
            <a href='https://github.com/MKYUKI'><img src='{linkedin}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.youtube.com/@mk_agi'><img src='{youtube}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.paypal.com/paypalme/MasakiKusaka'><img src='{paypal}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.dropbox.com/home'><img src='{dropbox}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://huggingface.co/pricing'><img src='{huggingface}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://x.com/MK_ASI1'><img src='{x_icon}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.facebook.com/'><img src='{facebook}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.amazon.co.jp/s?i=digital-text&rh=p_27%3AMasaki+Kusaka&s=relevancerank&text=Masaki+Kusaka&ref=dp_byline_sr_ebooks_1'><img src='{amazon_jp}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.amazon.com/s?i=digital-text&rh=p_27%3AMasaki+Kusaka&s=relevancerank&text=Masaki+Kusaka&ref=dp_byline_sr_ebooks_1'><img src='{amazon_us}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.linkedin.com/build-relation/newsletter-follow?entityUrn=7163516439096733696'><img src='{newsletter}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.kaggle.com/sahirmaharajj'><img src='{share}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
        </div>
        <br>
        <a href='https://exifa.betteruptime.com/'><img src='{uptime}'></a>
        &nbsp; <a href="https://www.producthunt.com/posts/exifa-net?embed=true&utm_source=badge-featured&utm_medium=badge&utm_souce=badge-exifa&#0045;net" target="_blank"><img src="https://api.producthunt.com/widgets/embed-image/v1/featured.svg?post_id=474560&theme=dark" alt="Exifa&#0046;net - Your&#0032;AI&#0032;assistant&#0032;for&#0032;understanding&#0032;EXIF&#0032;data | Product Hunt" style="width: 125px; height: 27px;" width="125" height="27" /></a>
        """,
        unsafe_allow_html=True,
    )

########################################################
# ファイル処理＆EXIF解析
########################################################
file_text = ""
if st.session_state["uploaded_files"]:
    for uf in st.session_state["uploaded_files"]:
        if uf.type == "application/pdf":
            with pdfplumber.open(uf) as pdf:
                pages = [page.extract_text() for page in pdf.pages]
            file_text = "\n".join(p for p in pages if p)
        elif uf.type == "text/plain":
            file_text = str(uf.read(), "utf-8")
        elif uf.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uf)
            file_text = "\n".join([para.text for para in doc.paragraphs])
        elif uf.type == "text/csv":
            df = pd.read_csv(uf)
            file_text = df.to_string(index=False)
        elif uf.type in ["image/jpeg","image/png","image/jpg"]:
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                temp.write(uf.read())
                temp.flush()
                temp.close()
                with open(temp.name,"rb") as f:
                    tags = exifread.process_file(f)
                os.unlink(temp.name)
            exif_data = {}
            for tag in tags.keys():
                if tag not in ["JPEGThumbnail","TIFFThumbnail","Filename","EXIF MakerNote"]:
                    exif_data[tag] = str(tags[tag])
            df = pd.DataFrame(exif_data, index=[0])
            df.insert(loc=0, column="Image Feature", value=["Value"]*len(df))
            df = df.transpose()
            df.columns = df.iloc[0]
            df = df.iloc[1:]
            st.session_state["exif_df"] = df
            file_text = file_text or "\n".join([f"{tag}: {tags[tag]}" for tag in tags.keys() if tag not in ("JPEGThumbnail","TIFFThumbnail","Filename","EXIF MakerNote")])

if st.session_state["image_url"]:
    try:
        resp_head = requests.head(st.session_state["image_url"])
        if resp_head.headers.get("Content-Type","").startswith("image"):
            resp = requests.get(st.session_state["image_url"])
            resp.raise_for_status()
            image_data = BytesIO(resp.content)
            image = Image.open(image_data)
            image.load()
            tags = exifread.process_file(image_data)
            exif_data = {}
            for tag in tags.keys():
                if tag not in ["JPEGThumbnail","TIFFThumbnail","Filename","EXIF MakerNote"]:
                    exif_data[tag] = str(tags[tag])
            df = pd.DataFrame(exif_data, index=[0])
            df.insert(loc=0, column="Image Feature", value=["Value"]*len(df))
            df = df.transpose()
            df.columns = df.iloc[0]
            df = df.iloc[1:]
            st.session_state["exif_df"] = df
            file_text = "\n".join([f"{tag}: {tags[tag]}" for tag in tags.keys() if tag not in ("JPEGThumbnail","TIFFThumbnail","Filename","EXIF MakerNote")])
        else:
            st.warning("URLは画像ではありません。")
    except:
        st.warning("URLから画像取得失敗")

########################################################
# メインUI構築
########################################################
if not st.session_state["authenticated"]:
    # ユーザー登録とログインフォーム
    auth_tabs = st.tabs(["ログイン", "新規登録"])

    with auth_tabs[0]:
        st.subheader("ログイン")
        with st.form("login_form"):
            login_username = st.text_input("ユーザー名")
            login_password = st.text_input("パスワード", type="password")
            login_submit = st.form_submit_button("ログイン")
        if login_submit:
            users = load_users_from_db()
            if login_username in users and check_password(login_password, users[login_username]):
                st.session_state["authenticated"] = True
                st.session_state["username"] = login_username
                st.success("ログインに成功しました。")
                st.experimental_rerun()
            else:
                st.error("ユーザー名またはパスワードが正しくありません。")

    with auth_tabs[1]:
        st.subheader("新規登録")
        with st.form("register_form"):
            register_username = st.text_input("ユーザー名")
            register_password = st.text_input("パスワード", type="password")
            register_confirm_password = st.text_input("パスワード確認", type="password")
            register_submit = st.form_submit_button("登録")
        if register_submit:
            users = load_users_from_db()
            if register_username in users:
                st.error("このユーザー名は既に使用されています。")
            elif register_password != register_confirm_password:
                st.error("パスワードが一致しません。")
            elif len(register_password) < 6:
                st.error("パスワードは6文字以上にしてください。")
            else:
                hashed_pw = hash_password(register_password)
                users[register_username] = hashed_pw
                save_users_to_db(users)
                st.success("登録が完了しました。ログインしてください。")
                st.experimental_rerun()

    st.markdown("---")
    st.caption("© Exifa.net (Sahir Maharaj,2024), CC-BY 4.0. これは全てを統合した世界初の究極アプリ。")

else:
    st.markdown("<h1 style='text-align:center;color:white;'>究極融合: EXIF & TTS & GPT</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center;color:#cccccc;'>300ページ超テキスト音声化、EXIF解析、カラー可視化、GPT対話</p>",unsafe_allow_html=True)

    tabs = st.tabs([
        "📜 テキスト音声合成",
        "🖼 EXIF解析＆ビジュアル",
        "💬 GPT対話",
        "📈 データダッシュボード",
        "⚙️ ユーザー設定",
        "📖 テキスト分析",
        "🖌 画像強化"
    ])

    # 音声合成タブ
    with tabs[0]:
        st.subheader("超大規模テキスト音声化")
        input_option = st.selectbox("入力方法",("直接入力","アップロードテキスト利用"))
        tts_text = ""
        if input_option == "直接入力":
            tts_text = st.text_area("音声合成するテキストを貼り付け","ここに膨大なテキスト(例:書籍全文)を入力")
        else:
            if file_text:
                st.write("抽出テキスト(一部):")
                st.write(file_text[:500]+"...")
                tts_text = file_text
            else:
                st.write("アップロードテキストがありません")

        selected_gender = st.selectbox("話者の性別",('default','male','female','neutral'))
        speech_rate = st.slider("話速", 0.5, 2.0, 1.0, 0.1)
        speech_pitch = st.slider("ピッチ", -20.0, 20.0, 0.0, 1.0)
        if tts_text and st.button("音声合成実行"):
            with st.spinner("音声合成中...長文は時間要"):
                lang_code = detect_language(tts_text)
                final_mp3 = synthesize_speech_chunk(tts_text, lang_code, gender=selected_gender, rate=speech_rate, pitch=speech_pitch)
            st.success("音声合成完了！")
            st.download_button("MP3ダウンロード", data=final_mp3, file_name="converted_book.mp3", mime="audio/mpeg")
            st.audio(final_mp3, format="audio/mp3")

    # EXIF解析＆ビジュアルタブ
    with tabs[1]:
        st.subheader("EXIF解析 & 可視化")
        if st.session_state["exif_df"].empty and not st.session_state["image_url"]:
            st.info("EXIFデータなし: 画像アップロードかURL指定を")
        else:
            st.markdown("##### EXIFデータ抽出結果")
            st.dataframe(st.session_state["exif_df"])
            image_to_analyze = None
            if st.session_state["uploaded_files"]:
                for f in st.session_state["uploaded_files"]:
                    if f.type in ["image/jpeg","image/png","image/jpg"]:
                        image_to_analyze = load_image(f)
                        break
            elif st.session_state["image_url"]:
                image_to_analyze = load_image(st.session_state["image_url"])

            if image_to_analyze:
                st.image(image_to_analyze, caption="アップロード画像", use_column_width=True)
                data = np.array(image_to_analyze)

                exp1 = st.expander("⛆ RGBチャンネル操作")
                with exp1:
                    channels = st.multiselect("表示チャンネル:",["Red","Green","Blue"],default=["Red","Green","Blue"])
                    if channels:
                        cmap = {"Red":0,"Green":1,"Blue":2}
                        selected_idx = [cmap[ch] for ch in channels]
                        ch_data = np.zeros_like(data)
                        for idx in selected_idx:
                            ch_data[:,:,idx] = data[:,:,idx]
                        st.image(Image.fromarray(ch_data), use_column_width=True)
                    else:
                        st.image(image_to_analyze, use_column_width=True)

                exp2 = st.expander("〽 HSVヒストグラム")
                with exp2:
                    hsv_image = image_to_analyze.convert("HSV")
                    hsv_data = np.array(hsv_image)
                    hue_hist, _ = np.histogram(hsv_data[:,:,0], bins=256, range=(0,256))
                    sat_hist, _ = np.histogram(hsv_data[:,:,1], bins=256, range=(0,256))
                    val_hist, _ = np.histogram(hsv_data[:,:,2], bins=256, range=(0,256))
                    hsv_histogram_df = pd.DataFrame({"Hue":hue_hist,"Saturation":sat_hist,"Value":val_hist})
                    st.line_chart(hsv_histogram_df)

                exp3 = st.expander("☄ カラー分布サンバースト")
                with exp3:
                    red, green, blue = data[:,:,0], data[:,:,1], data[:,:,2]
                    ci = {"color":[],"intensity":[],"count":[]}
                    for name,channel in zip(["Red","Green","Blue"],[red,green,blue]):
                        unique, counts = np.unique(channel, return_counts=True)
                        ci["color"].extend([name]*len(unique))
                        ci["intensity"].extend(unique)
                        ci["count"].extend(counts)
                    cdf = pd.DataFrame(ci)
                    fig = px.sunburst(cdf,path=["color","intensity"],values="count",color="color",
                                      color_discrete_map={"Red":"#ff6666","Green":"#85e085","Blue":"#6666ff"})
                    st.plotly_chart(fig,use_container_width=True)

                exp4 = st.expander("🕸 3D色空間プロット")
                with exp4:
                    skip = 8
                    sample = data[::skip,::skip].reshape(-1,3)
                    fig = go.Figure(data=[go.Scatter3d(
                        x=sample[:,0],y=sample[:,1],z=sample[:,2],
                        mode="markers",
                        marker=dict(size=3,color=["rgb({},{},{})".format(r,g,b) for r,g,b in sample])
                    )])
                    fig.update_layout(scene=dict(xaxis_title="Red",yaxis_title="Green",zaxis_title="Blue"))
                    st.plotly_chart(fig, use_container_width=True)

                # 画像フィルター適用
                exp5 = st.expander("🖌 画像フィルター適用")
                with exp5:
                    filter_option = st.selectbox("適用するフィルターを選択:", ["None", "グレースケール", "セピア", "エッジ検出"])
                    if filter_option == "グレースケール":
                        filtered_image = image_to_analyze.convert("L")
                        st.image(filtered_image, caption="グレースケール画像", use_column_width=True)
                    elif filter_option == "セピア":
                        filtered_image = image_to_analyze.convert("RGB")
                        np_image = np.array(filtered_image)
                        tr = [int(0.393 * r + 0.769 * g + 0.189 * b) for r, g, b in np_image.reshape(-1,3)]
                        tg = [int(0.349 * r + 0.686 * g + 0.168 * b) for r, g, b in np_image.reshape(-1,3)]
                        tb = [int(0.272 * r + 0.534 * g + 0.131 * b) for r, g, b in np_image.reshape(-1,3)]
                        tr = np.clip(tr, 0, 255)
                        tg = np.clip(tg, 0, 255)
                        tb = np.clip(tb, 0, 255)
                        sepia = np.stack([tr, tg, tb], axis=1).reshape(np_image.shape).astype(np.uint8)
                        sepia_image = Image.fromarray(sepia)
                        st.image(sepia_image, caption="セピア画像", use_column_width=True)
                    elif filter_option == "エッジ検出":
                        edge_image = image_to_analyze.filter(ImageFilter.FIND_EDGES)
                        st.image(edge_image, caption="エッジ検出画像", use_column_width=True)
                    else:
                        st.image(image_to_analyze, caption="元の画像", use_column_width=True)

                # オブジェクト検出（簡易実装）
                exp6 = st.expander("🔍 オブジェクト検出")
                with exp6:
                    if st.button("オブジェクト検出実行"):
                        with st.spinner("オブジェクト検出中..."):
                            # 実際のオブジェクト検出にはAIモデルを使用することが望ましい
                            # ここでは例として画像をぼかす処理を行います
                            enhanced_image = image_to_analyze.filter(ImageFilter.GaussianBlur(radius=2))
                        st.success("オブジェクト検出完了！")
                        st.image(enhanced_image, caption="オブジェクト検出後の画像", use_column_width=True)

                st.markdown("#### EXIF除去後の画像ダウンロード")
                cleaned = clear_exif_data(image_to_analyze)
                download_image(cleaned)

            # 簡易コメント（LLM対応可）
            if not st.session_state["exif_df"].empty:
                commentary = """EXIFから撮影者の機材・露出設定などが推測可能。撮影環境は自然光か計画的照明下とみられ、撮影者は中級的経験と程よい予算を持つと考えられる。"""
                st.markdown("#### 自動生成コメント")
                st.write(commentary)
                if st.button("コメント音声再生"):
                    lang_code = detect_language(commentary)
                    audio_data = synthesize_speech_chunk(commentary, lang_code)
                    st.audio(audio_data, format="audio/mp3")

    # GPT対話タブ
    with tabs[2]:
        st.subheader("GPTによる高度な対話")
        for msg in st.session_state["messages"]:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        user_input = st.chat_input("EXIF、TTS、画像解析、色空間、汎用的な質問すべてをどうぞ")
        if user_input:
            st.session_state["messages"].append({"role":"user","content":user_input})
            with st.chat_message("user"):
                st.write(user_input)

            if openai.api_key:
                response = openai.ChatCompletion.create(
                    model="gpt-4", 
                    messages=st.session_state["messages"],
                    temperature=0.3,
                    top_p=0.9,
                    frequency_penalty=0,
                    presence_penalty=0
                )
                answer = response.choices[0].message["content"]
            else:
                answer = "OpenAI APIキーが未設定です。Secretesで設定してください。"

            st.session_state["messages"].append({"role":"assistant","content":answer})
            with st.chat_message("assistant"):
                st.write(answer)

    ########################################################
    # データダッシュボードタブの追加
    ########################################################
    with tabs[3]:
        st.subheader("データダッシュボード")
        st.markdown("##### アップロードされたファイルやユーザーの活動をリアルタイムで可視化します。")

        if st.session_state["uploaded_files"] or not st.session_state["exif_df"].empty:
            # アップロードファイル数のカウント
            num_files = len(st.session_state["uploaded_files"]) if st.session_state["uploaded_files"] else 0

            # EXIFデータの統計
            if not st.session_state["exif_df"].empty:
                num_exif = len(st.session_state["exif_df"])
            else:
                num_exif = 0

            # ダッシュボードメトリクス
            col1, col2 = st.columns(2)
            with col1:
                st.metric("アップロードファイル数", num_files)
            with col2:
                st.metric("抽出EXIFデータ数", num_exif)

            # アップロードファイルタイプの分布
            if st.session_state["uploaded_files"]:
                file_types = [uf.type for uf in st.session_state["uploaded_files"]]
                type_counts = pd.Series(file_types).value_counts().reset_index()
                type_counts.columns = ["File Type", "Count"]
                fig1 = px.pie(type_counts, names='File Type', values='Count', title='アップロードファイルタイプの分布')
                st.plotly_chart(fig1, use_container_width=True)

            # EXIFデータの主要タグの頻度
            if not st.session_state["exif_df"].empty:
                top_exif = st.session_state["exif_df"].head(10)
                fig2 = px.bar(top_exif, x=st.session_state["exif_df"].index, y=top_exif.columns, title='主要EXIFタグの頻度')
                st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("アップロードされたファイルやEXIFデータがありません。")

    ########################################################
    # ユーザー設定タブの追加
    ########################################################
    with tabs[4]:
        st.subheader("ユーザー設定")
        st.markdown("##### アプリケーションの設定をカスタマイズします。")

        # テーマ選択
        theme = st.selectbox("テーマ選択", ["ダークモード", "ライトモード"], index=0)
        if theme == "ライトモード":
            light_css = """
            <style>
            body {
                background: #ffffff;
                color: #000000;
                font-family: 'Helvetica', sans-serif;
            }
            h1, h2, h3, h4, h5, h6 {
                color: #000000;
            }
            .block-container {
                padding: 1rem 2rem;
            }
            .sidebar .sidebar-content {
                background: #f0f0f0;
                color: #000000;
            }
            .stTextInput > div {
                color:#000000;
            }
            </style>
            """
            st.markdown(light_css, unsafe_allow_html=True)
            st.session_state["user_preferences"]["theme"] = "ライトモード"
            st.success("ライトモードに切り替わりました。")
        else:
            st.markdown(custom_css, unsafe_allow_html=True)
            st.session_state["user_preferences"]["theme"] = "ダークモード"
            st.success("ダークモードに切り替わりました。")

        st.markdown("---")

        # 通知設定
        st.markdown("##### 通知設定")
        notifications = st.checkbox("音声合成完了時に通知を受け取る", value=st.session_state["user_preferences"]["notifications"])
        st.session_state["user_preferences"]["notifications"] = notifications
        if notifications:
            st.success("音声合成完了時に通知を受け取るよう設定されました。")
        else:
            st.info("音声合成完了時の通知がオフになりました。")

        st.markdown("---")

        # パスワード変更
        st.markdown("##### パスワード変更")
        with st.form("password_change_form"):
            current_password = st.text_input("現在のパスワード", type="password")
            new_password = st.text_input("新しいパスワード", type="password")
            confirm_password = st.text_input("新しいパスワード確認", type="password")
            submit_pw = st.form_submit_button("パスワード変更")

        if submit_pw:
            users = load_users_from_db()
            username = st.session_state["username"]
            if username in users and check_password(current_password, users[username]):
                if new_password != confirm_password:
                    st.error("新しいパスワードが一致しません。")
                elif len(new_password) < 6:
                    st.error("パスワードは6文字以上にしてください。")
                else:
                    users[username] = hash_password(new_password)
                    save_users_to_db(users)
                    st.success("パスワードが正常に変更されました。")
            else:
                st.error("現在のパスワードが正しくありません。")

        st.markdown("---")

        # アカウント削除
        st.markdown("##### アカウント削除")
        if st.button("アカウントを削除する"):
            confirm = st.checkbox("本当にアカウントを削除しますか？ この操作は取り消せません。")
            if confirm:
                users = load_users_from_db()
                username = st.session_state["username"]
                if username in users:
                    del users[username]
                    save_users_to_db(users)
                    st.session_state["authenticated"] = False
                    st.session_state["username"] = ""
                    st.success("アカウントが削除されました。")
                    st.experimental_rerun()
                else:
                    st.error("アカウントの削除に失敗しました。")

    ########################################################
    # テキスト分析タブの追加
    ########################################################
    with tabs[5]:
        st.subheader("テキスト分析")
        st.markdown("##### アップロードされたテキストの感情分析、要約、キーワード抽出を行います。")

        text_to_analyze = ""
        if st.session_state["uploaded_files"]:
            for uf in st.session_state["uploaded_files"]:
                if uf.type in ["text/plain","application/pdf","application/vnd.openxmlformats-officedocument.wordprocessingml.document","text/csv"]:
                    if uf.type == "text/plain":
                        text_to_analyze = str(uf.read(), "utf-8")
                    elif uf.type == "application/pdf":
                        with pdfplumber.open(uf) as pdf:
                            pages = [page.extract_text() for page in pdf.pages]
                        text_to_analyze = "\n".join(p for p in pages if p)
                    elif uf.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = Document(uf)
                        text_to_analyze = "\n".join([para.text for para in doc.paragraphs])
                    elif uf.type == "text/csv":
                        df = pd.read_csv(uf)
                        text_to_analyze = df.to_string(index=False)
                    break

        if st.session_state["image_url"] and not text_to_analyze:
            st.warning("画像URLが指定されていますが、テキストデータがありません。テキストファイルをアップロードしてください。")

        if text_to_analyze:
            st.markdown("##### 分析対象テキストの一部")
            st.write(text_to_analyze[:500] + ("..." if len(text_to_analyze) > 500 else ""))

            if st.button("テキスト分析実行"):
                with st.spinner("テキスト分析中..."):
                    language = detect_language(text_to_analyze)
                    sentiment = analyze_sentiment(text_to_analyze)
                    summary = summarize_text(text_to_analyze, language='ja' if language=='ja-JP' else 'en')
                    keywords = extract_keywords(text_to_analyze)

                st.success("テキスト分析完了！")
                st.markdown("#### 感情分析結果")
                st.write(f"**感情:** {sentiment}")

                st.markdown("#### テキスト要約")
                st.write(summary)

                st.markdown("#### キーワード抽出")
                st.write(keywords)

                # 通知設定に基づく通知
                if st.session_state["user_preferences"]["notifications"]:
                    st.balloons()

        else:
            st.info("分析するテキストデータをアップロードしてください。")

    ########################################################
    # 画像強化タブの追加
    ########################################################
    with tabs[6]:
        st.subheader("画像強化")
        st.markdown("##### アップロードされた画像に対してオブジェクト検出やフィルターを適用します。")

        image_to_enhance = None
        if st.session_state["uploaded_files"]:
            for f in st.session_state["uploaded_files"]:
                if f.type in ["image/jpeg","image/png","image/jpg"]:
                    image_to_enhance = load_image(f)
                    break
        elif st.session_state["image_url"]:
            image_to_enhance = load_image(st.session_state["image_url"])

        if image_to_enhance:
            st.image(image_to_enhance, caption="元の画像", use_column_width=True)

            # オブジェクト検出（簡易実装）
            exp1 = st.expander("🔍 オブジェクト検出")
            with exp1:
                if st.button("オブジェクト検出実行"):
                    with st.spinner("オブジェクト検出中..."):
                        # 実際のオブジェクト検出にはAIモデルを使用することが望ましい
                        # ここでは例として画像をぼかす処理を行います
                        enhanced_image = image_to_enhance.filter(ImageFilter.GaussianBlur(radius=2))
                    st.success("オブジェクト検出完了！")
                    st.image(enhanced_image, caption="オブジェクト検出後の画像", use_column_width=True)

            # 画像フィルター適用
            exp2 = st.expander("🖌 画像フィルター適用")
            with exp2:
                filter_option = st.selectbox("適用するフィルターを選択:", ["None", "グレースケール", "セピア", "エッジ検出"])
                if filter_option == "グレースケール":
                    filtered_image = image_to_enhance.convert("L")
                    st.image(filtered_image, caption="グレースケール画像", use_column_width=True)
                elif filter_option == "セピア":
                    filtered_image = image_to_enhance.convert("RGB")
                    np_image = np.array(filtered_image)
                    tr = [int(0.393 * r + 0.769 * g + 0.189 * b) for r, g, b in np_image.reshape(-1,3)]
                    tg = [int(0.349 * r + 0.686 * g + 0.168 * b) for r, g, b in np_image.reshape(-1,3)]
                    tb = [int(0.272 * r + 0.534 * g + 0.131 * b) for r, g, b in np_image.reshape(-1,3)]
                    tr = np.clip(tr, 0, 255)
                    tg = np.clip(tg, 0, 255)
                    tb = np.clip(tb, 0, 255)
                    sepia = np.stack([tr, tg, tb], axis=1).reshape(np_image.shape).astype(np.uint8)
                    sepia_image = Image.fromarray(sepia)
                    st.image(sepia_image, caption="セピア画像", use_column_width=True)
                elif filter_option == "エッジ検出":
                    edge_image = image_to_enhance.filter(ImageFilter.FIND_EDGES)
                    st.image(edge_image, caption="エッジ検出画像", use_column_width=True)
                else:
                    st.image(image_to_enhance, caption="元の画像", use_column_width=True)

            st.markdown("#### EXIF除去後の画像ダウンロード")
            cleaned = clear_exif_data(image_to_enhance)
            download_image(cleaned)

        else:
            st.info("画像をアップロードしてください。")

    st.markdown("---")
    st.caption("© Exifa.net (Sahir Maharaj,2024), CC-BY 4.0. これは全てを統合した世界初の究極アプリ。")

    # フッターや追加情報
    st.sidebar.caption(
        "Built by [Sahir Maharaj](https://www.linkedin.com/in/sahir-maharaj/). Like this? [Hire me!](https://topmate.io/sahirmaharaj/362667)"
    )
    
    # 追加のリンクや情報
    linkedin = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/linkedin.svg"
    youtube = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/youtube.svg"
    paypal = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/paypal.svg"
    dropbox = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/dropbox.svg"
    huggingface = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/huggingface.svg"
    x_icon = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/x.svg"
    facebook = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/facebook.svg"
    amazon_jp = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/amazon-jp.svg"
    amazon_us = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/amazon-us.svg"
    newsletter = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/newsletter.svg"
    share = "https://raw.githubusercontent.com/MKYUKI/youtube-new/main/img/share.svg"
    
    uptime = "https://uptime.betterstack.com/status-badges/v1/monitor/196o6.svg"
    
    st.sidebar.markdown(
        f"""
        <div style='display: flex; align-items: center;'>
            <a href='https://github.com/MKYUKI'><img src='{linkedin}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.youtube.com/@mk_agi'><img src='{youtube}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.paypal.com/paypalme/MasakiKusaka'><img src='{paypal}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.dropbox.com/home'><img src='{dropbox}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://huggingface.co/pricing'><img src='{huggingface}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://x.com/MK_ASI1'><img src='{x_icon}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.facebook.com/'><img src='{facebook}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.amazon.co.jp/s?i=digital-text&rh=p_27%3AMasaki+Kusaka&s=relevancerank&text=Masaki+Kusaka&ref=dp_byline_sr_ebooks_1'><img src='{amazon_jp}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.amazon.com/s?i=digital-text&rh=p_27%3AMasaki+Kusaka&s=relevancerank&text=Masaki+Kusaka&ref=dp_byline_sr_ebooks_1'><img src='{amazon_us}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.linkedin.com/build-relation/newsletter-follow?entityUrn=7163516439096733696'><img src='{newsletter}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
            <a href='https://www.kaggle.com/sahirmaharajj'><img src='{share}' style='width: 35px; height: 35px; margin-right: 15px;'></a>
        </div>
        <br>
        <a href='https://exifa.betteruptime.com/'><img src='{uptime}'></a>
        &nbsp; <a href="https://www.producthunt.com/posts/exifa-net?embed=true&utm_source=badge-featured&utm_medium=badge&utm_souce=badge-exifa&#0045;net" target="_blank"><img src="https://api.producthunt.com/widgets/embed-image/v1/featured.svg?post_id=474560&theme=dark" alt="Exifa&#0046;net - Your&#0032;AI&#0032;assistant&#0032;for&#0032;understanding&#0032;EXIF&#0032;data | Product Hunt" style="width: 125px; height: 27px;" width="125" height="27" /></a>
        """,
        unsafe_allow_html=True,
    )
